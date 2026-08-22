from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor


def build_docx(title: str, content: dict[str, Any]) -> bytes:
    doc = Document()

    doc.add_heading(title, level=1)

    case = content["case"]
    meta = doc.add_paragraph()
    meta.add_run(f"Справа: {case['title']}").italic = True
    if case.get("goal"):
        meta.add_run(f" · {case['goal']}").italic = True

    primary = content.get("primary")
    if primary:
        doc.add_heading(primary["display_name"], level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for f in primary["fields"]:
            row = table.add_row().cells
            label = f["label"] + (" *" if f["required"] else "")
            row[0].text = label
            row[1].text = str(f["value"]) if f["value"] else "—"
            if f["required"] and not f["value"]:
                for cell in row:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    for section in content["related_sections"]:
        if not section["entries"]:
            continue
        doc.add_heading(f"{section['label']} ({len(section['entries'])})", level=3)
        for item in section["entries"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item["display_name"])
            run.bold = True
            if item.get("relationship_type"):
                p.add_run(f" — {item['relationship_type']}")
            p.add_run(f" ({item['confidence']})")
            if item.get("details"):
                detail_line = "; ".join(f"{k}: {v}" for k, v in item["details"].items())
                doc.add_paragraph(detail_line, style="List Bullet 2")

    if content["events"]:
        doc.add_heading("Таймлайн", level=3)
        for e in content["events"]:
            line = f"{e['date']} — {e['title']}" if e.get("date") else e["title"]
            if e.get("description"):
                line += f": {e['description']}"
            doc.add_paragraph(line, style="List Bullet")

    if content["missing"]:
        doc.add_heading("Не вистачає даних", level=3)
        for m in content["missing"]:
            p = doc.add_paragraph(m, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    footer = doc.add_paragraph()
    footer_run = footer.add_run("Згенеровано OSINT HUB · чернетка потребує перевірки аналітиком")
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
